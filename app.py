import streamlit as st
import streamlit.components.v1 as components
import hmac
import io
import re
import os
import requests
from datetime import datetime
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from src.extractor import process_file, sanitize_filename
from src.qa_engine import QAEngine
from src.config import validate_config, APP_PASSWORD, LOG_LEVEL
from src.logging_utils import setup_logging
import time
from src.mermaid_renderer import MermaidCleaner, MermaidValidator, MERMAID_HTML_TEMPLATE, MERMAID_CDN, render_content_with_mermaid

# Initialize logging at application startup with configured level
setup_logging(level=LOG_LEVEL)

# Maximum chat input length to prevent abuse
MAX_CHAT_INPUT_LENGTH = 2000


# sanitize_filename is imported from src.extractor to avoid duplication


def sanitize_for_markdown(text: str) -> str:
    """Sanitize user-provided text to prevent HTML/Markdown injection."""
    if not text:
        return ""
    # Escape HTML special characters
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    # Escape markdown special characters that could break rendering
    text = text.replace("[", "&#91;").replace("]", "&#93;")
    return text





def render_confidence_badge(confidence: dict) -> None:
    """
    Renders a colour-coded confidence percentage badge with 4-factor breakdown.
    Displayed inside the assistant message bubble, after the answer.
    """
    score  = confidence.get("score", 50)
    level  = confidence.get("level", "Moderate")
    reason = confidence.get("reason", "")

    colour_map = {
        "Very High": "#22c55e",
        "High":      "#3b82f6",
        "Moderate":  "#f59e0b",
        "Low":       "#f97316",
        "Very Low":  "#ef4444",
    }
    colour = colour_map.get(level, "#94a3b8")

    badge_html = f"""
    <div>
        <div class="confidence-wrapper">
            <span
                class="confidence-badge"
                style="background-color: {colour};"
                title="{level} — {reason}"
            >
                <strong>{score}%</strong>
            </span>
            <span class="confidence-label">AI Confidence</span>
        </div>
    </div>
    """
    st.markdown(badge_html, unsafe_allow_html=True)


def strip_image_prompts(text: str) -> str:  # noqa: D401
    """Remove <image_prompt>...</image_prompt> tags from content so images don't appear in chat."""
    if not text:
        return text
    return re.sub(r"<image_prompt>.*?</image_prompt>", "", text, flags=re.DOTALL).strip()


def _clean_for_speech(text):
    text = re.sub(
        r"<image_prompt>.*?</image_prompt>",
        "Image generated.",
        text,
        flags=re.DOTALL,
    )
    text = re.sub(
        r"```mermaid.*?```", "Diagram included.", text, flags=re.DOTALL
    )
    text = re.sub(r"```.*?```", "", text, flags=re.DOTALL)
    text = re.sub(r"#{1,6}\s*", "", text)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"`(.*?)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"[|\-]{3,}", "", text)
    text = re.sub(r"<[^>]+>", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def render_mic_input() -> None:
    """Renders a compact mic input button."""
    if st.button("🎙️", help="Click to speak your question", key="mic_btn_compact"):
        try:
            import speech_recognition as sr
        except ImportError:
            st.warning("🎙️ Voice input requires: `pip install SpeechRecognition PyAudio`")
            return

        try:
            recognizer = sr.Recognizer()
            recognizer.energy_threshold = 300
            recognizer.dynamic_energy_threshold = True
            with sr.Microphone() as source:
                with st.spinner("🎤 Listening..."):
                    recognizer.adjust_for_ambient_noise(source, duration=0.5)
                    audio = recognizer.listen(source, timeout=10, phrase_time_limit=15)
            with st.spinner("📝..."):
                text = recognizer.recognize_google(audio)
                if text:
                    st.session_state["mic_text"] = text
                    st.rerun()
        except Exception as e:
            st.warning(f"🎙️ Error: {str(e)}")


def speak_text(text: str, key: str) -> None:
    clean = _clean_for_speech(text)
    safe = (
        clean.replace("\\", "\\\\")
        .replace("`", "")
        .replace('"', "&quot;")
        .replace("'", "\\'")
        .replace("\n", " ")
        .replace("\r", "")
    )
    # Truncate for long content
    if len(safe) > 5000:
        safe = safe[:5000] + "... Content truncated."
    
    html = f"""
    <div style="margin: 4px 0;">
        <button id="tts-btn-{key}" onclick="toggleSpeech_{key}()" style="
            background: linear-gradient(135deg, #302b63, #24243e);
            color: #a8edea; border: 1px solid #4a45a0; border-radius: 8px;
            padding: 6px 16px; cursor: pointer; font-size: 0.8rem;
            font-weight: 600; font-family: Inter, sans-serif;
            transition: all 0.2s ease; display: inline-flex; align-items: center; gap: 8px;
        ">
            <span id="tts-icon-{key}">🔊</span> <span id="tts-label-{key}">Listen</span>
        </button>
    </div>
    <script>
        var synth_{key} = window.speechSynthesis;
        var speaking_{key} = false;
        var utterance_{key} = null;

        function toggleSpeech_{key}() {{
            var btn = document.getElementById('tts-btn-{key}');
            var icon = document.getElementById('tts-icon-{key}');
            var label = document.getElementById('tts-label-{key}');

            if (speaking_{key}) {{
                synth_{key}.cancel();
                speaking_{key} = false;
                label.innerHTML = 'Listen';
                icon.innerHTML = '🔊';
                btn.style.borderColor = '#4a45a0';
            }} else {{
                synth_{key}.cancel(); // Interrupt any ongoing speech
                utterance_{key} = new SpeechSynthesisUtterance("{safe}");
                utterance_{key}.rate = 1.0;
                utterance_{key}.pitch = 1.0;
                
                utterance_{key}.onend = function() {{
                    speaking_{key} = false;
                    label.innerHTML = 'Listen';
                    icon.innerHTML = '🔊';
                    btn.style.borderColor = '#4a45a0';
                }};
                
                utterance_{key}.onerror = function(event) {{
                    console.error('TTS Error:', event);
                    speaking_{key} = false;
                    label.innerHTML = 'Error';
                    icon.innerHTML = '⚠️';
                }};

                synth_{key}.speak(utterance_{key});
                speaking_{key} = true;
                label.innerHTML = 'Stop';
                icon.innerHTML = '⏹️';
                btn.style.borderColor = '#a8edea';
            }}
        }}
    </script>
    """
    components.html(html, height=45)


def render_ollama_status_sidebar() -> None:
    """Queries local Ollama tags API to render a beautiful real-time status card in the sidebar."""
    from src.config import OLLAMA_BASE_URL, OLLAMA_PRIMARY_MODEL, OLLAMA_CODING_MODEL, OLLAMA_REASONING_MODEL
    
    connected = False
    models_status = {
        OLLAMA_PRIMARY_MODEL: "Missing",
        OLLAMA_CODING_MODEL: "Missing",
        OLLAMA_REASONING_MODEL: "Missing"
    }
    
    try:
        resp = requests.get(f"{OLLAMA_BASE_URL}/api/tags", timeout=1.5)
        if resp.status_code == 200:
            connected = True
            available_models = [m.get("name", "") for m in resp.json().get("models", [])]
            for req_model in models_status.keys():
                found = any(
                    m == req_model or m.startswith(req_model.split(":")[0] + ":")
                    for m in available_models
                )
                if found:
                    models_status[req_model] = "Loaded"
    except Exception:
        pass

    dot_class = "status-connected" if connected else "status-disconnected"
    status_label = "Online" if connected else "Offline"
    status_color = "#22c55e" if connected else "#ef4444"
    
    status_html = f"""
    <div class="status-card">
        <div class="status-header">
            <span>⚙️ Local Ollama Engine</span>
            <span style="display: flex; align-items: center; gap: 6px;">
                <span class="status-dot {dot_class}"></span>
                <span style="font-size: 0.75rem; font-weight: 600; color: {status_color};">{status_label}</span>
            </span>
        </div>
        <div style="border-top: 1px solid rgba(74, 69, 160, 0.2); margin-top: 6px; padding-top: 6px;">
            <div class="model-row">
                <span>🦙 General: <code>{OLLAMA_PRIMARY_MODEL}</code></span>
                <span class="model-badge" style="background: {'rgba(34, 197, 94, 0.15)' if models_status[OLLAMA_PRIMARY_MODEL] == 'Loaded' else 'rgba(239, 68, 68, 0.12)'}; color: {'#15803d' if models_status[OLLAMA_PRIMARY_MODEL] == 'Loaded' else '#b91c1c'} !important;">
                    {models_status[OLLAMA_PRIMARY_MODEL]}
                </span>
            </div>
            <div class="model-row">
                <span>🐉 Coding/Mermaid: <code>{OLLAMA_CODING_MODEL}</code></span>
                <span class="model-badge" style="background: {'rgba(34, 197, 94, 0.15)' if models_status[OLLAMA_CODING_MODEL] == 'Loaded' else 'rgba(239, 68, 68, 0.12)'}; color: {'#15803d' if models_status[OLLAMA_CODING_MODEL] == 'Loaded' else '#b91c1c'} !important;">
                    {models_status[OLLAMA_CODING_MODEL]}
                </span>
            </div>
            <div class="model-row">
                <span>🧪 Reasoning: <code>{OLLAMA_REASONING_MODEL}</code></span>
                <span class="model-badge" style="background: {'rgba(34, 197, 94, 0.15)' if models_status[OLLAMA_REASONING_MODEL] == 'Loaded' else 'rgba(239, 68, 68, 0.12)'}; color: {'#15803d' if models_status[OLLAMA_REASONING_MODEL] == 'Loaded' else '#b91c1c'} !important;">
                    {models_status[OLLAMA_REASONING_MODEL]}
                </span>
            </div>
        </div>
    </div>
    """
    st.markdown(status_html, unsafe_allow_html=True)



st.set_page_config(
    page_title="AI Powered Document Q&A System",
    page_icon="🔬",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Validate configuration on startup (cached to avoid blocking every rerun)
@st.cache_data(ttl=60, show_spinner=False)
def _cached_validate_config():
    return validate_config()

is_valid, config_errors = _cached_validate_config()
if not is_valid:
    st.error("⚠️ Configuration Error:")
    for error in config_errors:
        st.error(f"  • {error}")
    st.info(
        "Please make sure Ollama is running (`ollama serve`) and the required model is pulled:\n"
        "• `ollama pull llama3.1:8b`"
    )
    st.stop()

# ==================== AUTHENTICATION & RATE LIMITING ====================

# Basic Password Protection
if APP_PASSWORD:
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False

    if not st.session_state.authenticated:
        st.markdown(
            """
            <div style="text-align: center; margin-top: 50px;">
                <h1>🔬 AI Powered Document Q&A</h1>
                <p style="color: #a8edea;">Please enter the designated password to access the system.</p>
            </div>
            """, 
            unsafe_allow_html=True
        )
        password_col1, password_col2, password_col3 = st.columns([1, 2, 1])
        with password_col2:
            st.warning("⚠️ **Security Notice:** The system runs on local Ollama LLMs. Access is restricted.", icon="🔒")
            pwd_input = st.text_input("Enter Password", type="password")
            if st.button("Access System", use_container_width=True):
                if hmac.compare_digest(pwd_input, APP_PASSWORD):
                    st.session_state.authenticated = True
                    st.rerun()
                else:
                    st.error("Incorrect password.")
        st.stop()


# Simple Rate Limiter
if "rate_limits" not in st.session_state:
    st.session_state.rate_limits = {
        "upload": [],
        "chat": [],
        "voice": []
    }

def check_rate_limit(action: str, max_requests: int, window_seconds: int = 60) -> bool:
    """Enforces rate limiting on a specific action for the current user session."""
    now = time.time()
    # Filter to only keep requests within the sliding window
    history = [t for t in st.session_state.rate_limits[action] if now - t < window_seconds]
    st.session_state.rate_limits[action] = history
    
    if len(history) >= max_requests:
        return False
        
    st.session_state.rate_limits[action].append(now)
    return True

CUSTOM_CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');

/* Smooth scrolling for the entire app */
html {
    scroll-behavior: smooth;
}

/* Fix for scroll breaking - ensure smooth scrolling */
[data-testid="stAppViewContainer"] {
    scroll-behavior: smooth;
    -webkit-overflow-scrolling: touch;
}

/* Prevent layout shifts during loading */
.stMarkdown, .stMarkdown p {
    overflow-wrap: break-word;
}

/* Fix for expanders - smooth expand/collapse */
.streamlit-expanderHeader {
    transition: background-color 0.2s ease;
    border-radius: 8px !important;
    background-color: #f1f5f9 !important;
    border: 1px solid #e2e8f0 !important;
}

/* Fix for chat message rendering */
[data-testid="chat-message-container"] {
    transition: opacity 0.1s ease;
}

/* Prevent horizontal scroll breaking */
div[data-testid="stMarkdownContainer"] {
    overflow-x: hidden;
}

/* Enhanced Justification Class */
.justified-text {
    text-align: justify !important;
    text-justify: inter-word;
    line-height: 1.8;
    color: #334155;
}

/* Fix for source cards scrolling */
.source-card {
    overflow-wrap: break-word;
    word-wrap: break-word;
}

/* App Background Override (Light Premium) */
[data-testid="stAppViewContainer"] {
    font-family: 'Inter', sans-serif;
    background-color: #f8fafc;
}

/* Sidebar Styling Override (Clean Light-Gray) */
[data-testid="stSidebar"] {
    background-color: #f1f5f9;
    border-right: 1px solid #e2e8f0;
    color: #1e293b;
}

/* Force all sidebar text, labels, and widgets to use dark readable colors */
[data-testid="stSidebar"] * {
    color: #1e293b;
}
[data-testid="stSidebar"] h1,
[data-testid="stSidebar"] h2,
[data-testid="stSidebar"] h3,
[data-testid="stSidebar"] h4 {
    color: #0f172a !important;
}
[data-testid="stSidebar"] p,
[data-testid="stSidebar"] span,
[data-testid="stSidebar"] label,
[data-testid="stSidebar"] div {
    color: #1e293b;
}
[data-testid="stSidebar"] .stMarkdown p {
    color: #334155 !important;
}
[data-testid="stSidebar"] small {
    color: #475569 !important;
}
[data-testid="stSidebar"] [data-testid="stFileUploaderDropzone"] {
    background-color: #ffffff;
    border-color: #cbd5e1;
    color: #334155;
}
[data-testid="stSidebar"] [data-testid="stFileUploaderDropzone"] span,
[data-testid="stSidebar"] [data-testid="stFileUploaderDropzone"] p,
[data-testid="stSidebar"] [data-testid="stFileUploaderDropzone"] small {
    color: #475569 !important;
}
[data-testid="stSidebar"] button {
    color: #ffffff !important;
    background-color: #4f46e5 !important;
    border-color: #4338ca !important;
}
[data-testid="stSidebar"] button:hover {
    background-color: #4338ca !important;
}
[data-testid="stSidebar"] hr {
    border-color: #cbd5e1 !important;
}

/* Voice Interaction UI — mic button inside chat input bar */
[data-testid="stBottom"] {
    position: relative;
}

/* Style the mic button container to sit inside the chat input */
.mic-overlay-btn {
    position: fixed;
    bottom: 16px;
    right: 80px;
    z-index: 9999999;
}

.mic-overlay-btn button {
    width: 40px !important;
    height: 40px !important;
    min-height: 40px !important;
    border-radius: 50% !important;
    background: linear-gradient(135deg, #4f46e5, #06b6d4) !important;
    color: white !important;
    border: 1px solid #818cf8 !important;
    padding: 0 !important;
    font-size: 1.15rem !important;
    cursor: pointer !important;
    box-shadow: 0 2px 10px rgba(79, 70, 229, 0.25) !important;
    transition: all 0.2s cubic-bezier(0.4, 0, 0.2, 1) !important;
    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
}

.mic-overlay-btn button:hover {
    transform: scale(1.12) !important;
    box-shadow: 0 4px 14px rgba(79, 70, 229, 0.4) !important;
    border-color: #a5b4fc !important;
}

.mic-overlay-btn button:active {
    transform: scale(0.93) !important;
}



@keyframes pulse-red {
    0% { box-shadow: 0 0 0 0 rgba(239, 68, 68, 0.4); }
    70% { box-shadow: 0 0 0 8px rgba(239, 68, 68, 0); }
    100% { box-shadow: 0 0 0 0 rgba(239, 68, 68, 0); }
}

/* Fix for st.tabs rendering and visibility */
.stTabs [data-baseweb="tab-list"] {
    gap: 8px;
    border-bottom: 2px solid #e2e8f0;
}
.stTabs [data-baseweb="tab"] {
    border-radius: 8px 8px 0 0;
    padding: 8px 16px;
    background-color: #e2e8f0;
    color: #475569;
    font-weight: 500;
    transition: all 0.2s ease;
}
.stTabs [aria-selected="true"] {
    background-color: #4f46e5 !important;
    color: #ffffff !important;
    font-weight: 600;
}

/* Main Header (Clean Premium Gradient) */
.main-header {
    background: linear-gradient(135deg, #e0e7ff 0%, #f3e8ff 100%);
    padding: 2rem;
    border-radius: 16px;
    margin-bottom: 1.5rem;
    color: #1e1b4b;
    text-align: center;
    border: 1px solid rgba(99, 102, 241, 0.15);
    box-shadow: 0 10px 25px -5px rgba(99, 102, 241, 0.1);
}
.main-header h1 {
    margin: 0; font-size: 1.85rem; font-weight: 800;
    background: linear-gradient(90deg, #4f46e5, #7c3aed);
    -webkit-background-clip: text; -webkit-text-fill-color: transparent;
}
.main-header p {
    margin: 0.5rem 0 0; font-size: 0.95rem; font-weight: 500; color: #4338ca; opacity: 0.95;
}

/* Intent Badges (Light Pastel Colors) */
.intent-badge {
    display: inline-flex; align-items: center; gap: 6px;
    padding: 4px 12px; border-radius: 20px; font-size: 0.76rem;
    font-weight: 700; margin-bottom: 8px;
    text-transform: uppercase;
}
.intent-document_qa { background: #e0f2fe; color: #0369a1; border: 1px solid #bae6fd; }
.intent-suggestion_request { background: #fef3c7; color: #b45309; border: 1px solid #fde68a; }
.intent-research_addon { background: #dcfce7; color: #15803d; border: 1px solid #bbf7d0; }
.intent-off_topic { background: #fee2e2; color: #b91c1c; border: 1px solid #fca5a5; }

/* Source Citation Card (Premium White Style) */
.source-card {
    background: #ffffff; border: 1px solid #e2e8f0; border-radius: 10px;
    padding: 12px 16px; margin-bottom: 8px; font-size: 0.82rem;
    box-shadow: 0 2px 5px rgba(0, 0, 0, 0.02);
}
.source-card .source-header {
    display: flex; justify-content: space-between; margin-bottom: 4px;
    align-items: center;
}
.source-card .source-file { color: #4f46e5; font-weight: 700; }
.source-card .source-topic { color: #6d28d9; font-size: 0.75rem; font-weight: 600; }
.source-card .source-score {
    background: #e0e7ff; color: #4f46e5; padding: 2px 8px;
    border-radius: 10px; font-size: 0.7rem; font-weight: 700;
}
.source-card .source-preview { color: #475569; font-size: 0.78rem; margin-top: 4px; line-height: 1.5; }

/* Stats Grid & Cards */
.stats-grid {
    display: grid; grid-template-columns: 1fr 1fr; gap: 8px; margin: 10px 0;
}
.stat-card {
    background: #ffffff;
    border: 1px solid #e2e8f0; border-radius: 10px;
    padding: 12px; text-align: center;
    box-shadow: 0 2px 5px rgba(0,0,0,0.02);
}
.stat-card .stat-value {
    font-size: 1.5rem; font-weight: 800;
    background: linear-gradient(90deg, #4f46e5, #7c3aed);
    -webkit-background-clip: text; -webkit-text-fill-color: transparent;
}
.stat-card .stat-label { font-size: 0.72rem; color: #64748b; font-weight: 600; margin-top: 2px; }

/* Suggestions Cards */
.suggestion-card {
    background: #ffffff; border: 1px solid #e2e8f0; border-radius: 10px;
    padding: 12px 16px; margin-bottom: 8px; cursor: pointer;
    transition: all 0.2s ease;
    box-shadow: 0 2px 5px rgba(0,0,0,0.02);
}
.suggestion-card:hover {
    border-color: #6366f1; transform: translateY(-1px);
    box-shadow: 0 6px 15px rgba(99, 102, 241, 0.07);
}
.suggestion-title { color: #0f172a; font-weight: 700; font-size: 0.85rem; }
.suggestion-desc { color: #475569; font-size: 0.78rem; margin-top: 4px; line-height: 1.5; }
.suggestion-cat {
    display: inline-block; padding: 2px 8px; border-radius: 10px;
    font-size: 0.68rem; font-weight: 700; margin-top: 6px;
    text-transform: uppercase;
}
.cat-improvement { background: #e0f2fe; color: #0369a1; }
.cat-innovation { background: #f3e8ff; color: #6d28d9; }
.cat-gap { background: #fef3c7; color: #b45309; }
.cat-research { background: #dcfce7; color: #15803d; }
.cat-optimization { background: #fef08a; color: #854d0e; }

/* Topic Chips */
.topic-chip {
    display: inline-block; background: #ffffff; border: 1px solid #e2e8f0;
    border-radius: 16px; padding: 4px 12px; margin: 3px 4px 3px 0;
    font-size: 0.75rem; color: #334155; font-weight: 500;
    box-shadow: 0 1px 3px rgba(0,0,0,0.01);
}

/* Download Action Buttons (Stripe Gradient) */
.stDownloadButton > button {
    background: linear-gradient(135deg, #4f46e5, #6366f1) !important;
    color: white !important; border: 1px solid #4f46e5 !important;
    border-radius: 10px !important; font-weight: 600 !important;
    transition: all 0.3s ease !important;
    box-shadow: 0 4px 10px rgba(79, 70, 229, 0.15) !important;
}
.stDownloadButton > button:hover {
    background: linear-gradient(135deg, #4338ca, #4f46e5) !important;
    box-shadow: 0 6px 15px rgba(79, 70, 229, 0.3) !important;
    transform: translateY(-1px);
}

/* Confidence Badge Styles */
.confidence-wrapper {
    display: flex;
    justify-content: flex-end;
    align-items: center;
    gap: 6px;
    margin-top: -8px;
    margin-bottom: 6px;
}

.confidence-badge {
    display: inline-block;
    padding: 2px 10px;
    border-radius: 999px;
    color: #ffffff;
    font-size: 0.85rem;
    font-weight: 800;
    letter-spacing: 0.03em;
    cursor: help;
    user-select: none;
    white-space: nowrap;
    box-shadow: 0 1px 3px rgba(0,0,0,0.12);
    transition: transform 0.15s ease, box-shadow 0.15s ease;
}

.confidence-badge:hover {
    transform: scale(1.07);
    box-shadow: 0 3px 8px rgba(0,0,0,0.16);
}

.confidence-label {
    font-size: 0.70rem;
    color: #64748b;
    font-style: italic;
    letter-spacing: 0.02em;
}

/* Sidebar Ollama Status Card styles (Light Mode Glassmorphic) */
.status-card {
    background: rgba(255, 255, 255, 0.9);
    border: 1px solid rgba(99, 102, 241, 0.25);
    border-radius: 12px;
    padding: 14px;
    margin-bottom: 20px;
    box-shadow: 0 4px 12px rgba(0, 0, 0, 0.05);
    backdrop-filter: blur(8px);
}
.status-header {
    display: flex;
    justify-content: space-between;
    align-items: center;
    font-size: 0.85rem;
    font-weight: 700;
    color: #3730a3 !important;
    margin-bottom: 8px;
}
.status-dot {
    width: 9px;
    height: 9px;
    border-radius: 50%;
    display: inline-block;
}
.status-connected {
    background-color: #22c55e;
    box-shadow: 0 0 8px #22c55e;
    animation: glow-green 2s infinite ease-in-out;
}
.status-disconnected {
    background-color: #ef4444;
    box-shadow: 0 0 8px #ef4444;
    animation: pulse-red 2s infinite ease-in-out;
}
.model-row {
    display: flex;
    justify-content: space-between;
    align-items: center;
    font-size: 0.76rem;
    margin-top: 6px;
    color: #334155 !important;
}
.model-row span {
    color: #334155 !important;
}
.model-row code {
    background: rgba(99, 102, 241, 0.1);
    padding: 1px 4px;
    border-radius: 4px;
    color: #4338ca !important;
    font-weight: 600;
}
.model-badge {
    padding: 2px 8px;
    border-radius: 6px;
    font-size: 0.65rem;
    font-weight: 700;
    text-transform: uppercase;
}

@keyframes glow-green {
    0% { box-shadow: 0 0 2px #22c55e; }
    50% { box-shadow: 0 0 10px #22c55e; }
    100% { box-shadow: 0 0 2px #22c55e; }
}

/* ============================================= */
/*  RADIO TABS — Premium Pill Style              */
/* ============================================= */
[data-testid="stRadio"] > div {
    display: flex !important;
    gap: 6px !important;
    background: #e2e8f0;
    border-radius: 12px;
    padding: 4px;
}
[data-testid="stRadio"] > div > label {
    flex: 1 !important;
    text-align: center !important;
    padding: 10px 20px !important;
    border-radius: 10px !important;
    font-weight: 600 !important;
    font-size: 0.9rem !important;
    color: #475569 !important;
    cursor: pointer !important;
    transition: all 0.25s ease !important;
    background: transparent !important;
    border: none !important;
    margin: 0 !important;
}
[data-testid="stRadio"] > div > label:hover {
    background: rgba(99, 102, 241, 0.08) !important;
    color: #4f46e5 !important;
}
[data-testid="stRadio"] > div > label[data-checked="true"],
[data-testid="stRadio"] > div > label:has(input:checked) {
    background: linear-gradient(135deg, #4f46e5, #6366f1) !important;
    color: #ffffff !important;
    box-shadow: 0 4px 12px rgba(79, 70, 229, 0.25) !important;
}
/* Hide the native radio circle */
[data-testid="stRadio"] input[type="radio"] {
    display: none !important;
}
/* Hide the radio label text "Navigation" */
[data-testid="stRadio"] > label {
    display: none !important;
}

/* ============================================= */
/*  CHAT INPUT BAR — Clean Focus Ring            */
/* ============================================= */
[data-testid="stChatInput"] {
    border: none !important;
    background: transparent !important;
}
[data-testid="stChatInput"] textarea,
[data-testid="stChatInput"] input {
    padding-right: 100px !important;
    border: 2px solid #e2e8f0 !important;
    background-color: #ffffff !important;
    color: #0f172a !important;
    border-radius: 12px !important;
    font-size: 0.95rem !important;
    font-family: 'Inter', sans-serif !important;
    box-shadow: 0 4px 12px rgba(0, 0, 0, 0.03) !important;
    transition: border-color 0.2s ease, box-shadow 0.2s ease !important;
}
[data-testid="stChatInput"] textarea:focus,
[data-testid="stChatInput"] input:focus {
    border-color: #6366f1 !important;
    box-shadow: 0 0 0 3px rgba(99, 102, 241, 0.15), 0 4px 12px rgba(0, 0, 0, 0.03) !important;
    outline: none !important;
}
/* Chat send button */
[data-testid="stChatInput"] button {
    background: linear-gradient(135deg, #4f46e5, #6366f1) !important;
    color: white !important;
    border: none !important;
    border-radius: 10px !important;
}

/* ============================================= */
/*  CHAT MESSAGES — Alignment & Spacing          */
/* ============================================= */
[data-testid="stChatMessage"] {
    padding: 16px 20px !important;
    margin-bottom: 12px !important;
    border-radius: 12px !important;
    border: 1px solid #e2e8f0 !important;
    background: #ffffff !important;
    box-shadow: 0 2px 8px rgba(0, 0, 0, 0.02) !important;
}
/* User messages — slight indigo tint */
[data-testid="stChatMessage"]:has([data-testid="chatAvatarIcon-user"]) {
    background: linear-gradient(135deg, #eef2ff, #f5f3ff) !important;
    border-color: #c7d2fe !important;
}
/* Assistant messages */
[data-testid="stChatMessage"]:has([data-testid="chatAvatarIcon-assistant"]) {
    background: #ffffff !important;
    border-color: #e2e8f0 !important;
}
/* Chat avatar sizing */
[data-testid="stChatMessage"] [data-testid^="chatAvatarIcon"] {
    width: 32px !important;
    height: 32px !important;
}
/* Message content text */
[data-testid="stChatMessage"] .stMarkdown p {
    color: #1e293b !important;
    font-size: 0.95rem !important;
    line-height: 1.7 !important;
}

/* ============================================= */
/*  GLOBAL TEXT & HEADING COLORS                 */
/* ============================================= */
[data-testid="stAppViewContainer"] h1,
[data-testid="stAppViewContainer"] h2,
[data-testid="stAppViewContainer"] h3 {
    color: #0f172a !important;
    font-family: 'Inter', sans-serif !important;
}
[data-testid="stAppViewContainer"] p {
    color: #334155;
}

/* ============================================= */
/*  EXPANDER STYLING                             */
/* ============================================= */
[data-testid="stExpander"] {
    border: 1px solid #e2e8f0 !important;
    border-radius: 10px !important;
    overflow: hidden;
    margin: 8px 0;
}
[data-testid="stExpander"] summary {
    background: #f8fafc !important;
    color: #1e293b !important;
    font-weight: 600 !important;
    padding: 10px 16px !important;
}
[data-testid="stExpander"] summary:hover {
    background: #f1f5f9 !important;
}

/* ============================================= */
/*  SPINNER / LOADING                            */
/* ============================================= */
.stSpinner > div {
    border-top-color: #6366f1 !important;
}

/* ============================================= */
/*  DOWNLOAD BUTTONS (Main Area — NOT sidebar)   */
/* ============================================= */
[data-testid="stAppViewContainer"] .stDownloadButton > button {
    background: linear-gradient(135deg, #4f46e5, #6366f1) !important;
    color: white !important;
    border: 1px solid #4f46e5 !important;
    border-radius: 10px !important;
    font-weight: 600 !important;
}

/* ============================================= */
/*  TTS BUTTON — keep its own dark style         */
/* ============================================= */
[data-testid="stChatMessage"] iframe {
    border: none !important;
}

/* ============================================= */
/*  DIVIDER                                      */
/* ============================================= */
[data-testid="stAppViewContainer"] hr {
    border-color: #e2e8f0 !important;
    margin: 16px 0 !important;
}

/* ============================================= */
/*  RESPONSIVE — Mobile & Tablet                 */
/* ============================================= */
@media (max-width: 768px) {
    .main-header {
        padding: 1.2rem !important;
        border-radius: 12px !important;
    }
    .main-header h1 {
        font-size: 1.3rem !important;
    }
    .main-header p {
        font-size: 0.82rem !important;
    }
    [data-testid="stRadio"] > div > label {
        padding: 8px 12px !important;
        font-size: 0.82rem !important;
    }
    [data-testid="stChatMessage"] {
        padding: 12px 14px !important;
    }
    .stats-grid {
        grid-template-columns: 1fr !important;
    }
    .mic-overlay-btn {
        right: 60px !important;
        bottom: 12px !important;
    }
    [data-testid="stChatInput"] textarea,
    [data-testid="stChatInput"] input {
        padding-right: 70px !important;
        font-size: 0.88rem !important;
    }
}

@media (max-width: 480px) {
    .main-header h1 {
        font-size: 1.1rem !important;
    }
    [data-testid="stRadio"] > div {
        flex-direction: column !important;
    }
    [data-testid="stRadio"] > div > label {
        text-align: center !important;
    }
}

</style>
"""

st.markdown(CUSTOM_CSS, unsafe_allow_html=True)

st.markdown(
    """
<div class="main-header">
    <h1>🔬 AI Powered Document Q&A System</h1>
    <p>Upload documents • Ask questions • Get research insights • Download analysis</p>
</div>
""",
    unsafe_allow_html=True,
)

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "topics_found" not in st.session_state:
    st.session_state.topics_found = {}
if "system_ready" not in st.session_state:
    st.session_state.system_ready = False
if "ai_engine" not in st.session_state:
    st.session_state.ai_engine = QAEngine()
if "auto_suggestions" not in st.session_state:
    st.session_state.auto_suggestions = []
if "doc_overview" not in st.session_state:
    st.session_state.doc_overview = ""
if "ieee_metadata" not in st.session_state:
    st.session_state.ieee_metadata = {
        "title": "",
        "authors": "",
        "emails": "",
        "colleges": "",
        "additional_notes": "",
    }
# Lock for concurrent upload protection
if "processing_lock" not in st.session_state:
    st.session_state.processing_lock = False


def generate_markdown_export(
    history, overview="", suggestions=None, stats=None
):
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    md = "# 🔬 AI Powered Document Q&A System Report\n\n"
    md += f"**Generated**: {timestamp}\n\n---\n\n"

    if stats:
        md += "## 📊 Session Analytics\n\n"
        md += "| Metric | Value |\n|---|---|\n"
        md += f"| Questions Asked | {stats.get('questions_asked', 0)} |\n"
        md += f"| Topics Explored | {stats.get('topics_accessed', 0)} |\n"
        md += f"| Sources Referenced | {stats.get('sources_used', 0)} |\n"
        md += (
            f"| Total Topics Available | {stats.get('total_topics', 0)} |\n\n"
        )

    if overview:
        md += f"{overview}\n\n---\n\n"

    if suggestions:
        md += "## 💡 AI-Generated Suggestions\n\n"
        for i, s in enumerate(suggestions, 1):
            md += f"### {i}. {s.get('title', 'Suggestion')}\n"
            md += f"{s.get('description', '')}\n"
            md += f"*Category: {s.get('category', 'general')}*\n\n"
        md += "---\n\n"

    md += "## 💬 Research Session Log\n\n"
    for msg in history:
        role = "👤 Researcher" if msg["role"] == "user" else "🧠 AI Analyst"
        md += f"### {role}\n"

        if msg.get("intent"):
            intent = msg["intent"]
            md += f"*Intent: {intent.get('emoji', '')} {intent.get('label', '')}*\n\n"

        md += f"{msg['content']}\n\n"

        if (
            msg.get("reasoning_details")
            and msg["reasoning_details"]
            != "Model processed logic internally (Invisible Reasoning Pipeline)."
        ):
            md += f"<details><summary>Logic Trace</summary>\n\n> {msg['reasoning_details']}\n\n</details>\n\n"

        if msg.get("sources"):
            md += "**Sources Referenced:**\n"
            for src in msg["sources"]:
                md += f"- [{src['file']}] {src['topic']} (relevance: {src['score']})\n"
            md += "\n"

        md += "---\n\n"

    return md


def generate_html_export(history, overview="", suggestions=None, stats=None):
    html = """<!DOCTYPE html>
<html><head><meta charset="utf-8">
<title>AI Powered Document Q&A System Report</title>
<style>
body { font-family: 'Segoe UI', sans-serif; max-width: 900px; margin: 0 auto; padding: 40px; background: #0f0c29; color: #e0e0f0; }
h1 { background: linear-gradient(90deg, #a8edea, #fed6e3); -webkit-background-clip: text; -webkit-text-fill-color: transparent; text-align: center; }
h2 { color: #a8edea; border-bottom: 1px solid #313147; padding-bottom: 8px; }
.msg { background: #1e1e2e; border: 1px solid #313147; border-radius: 12px; padding: 16px 20px; margin: 12px 0; }
.user { border-left: 3px solid #a8edea; }
.assistant { border-left: 3px solid #fed6e3; }
.role { font-weight: 700; font-size: 0.9rem; margin-bottom: 8px; }
.intent-tag { display: inline-block; padding: 2px 10px; border-radius: 12px; font-size: 0.75rem; background: #302b63; color: #a8edea; margin-bottom: 8px; }
.stat-grid { display: grid; grid-template-columns: repeat(4, 1fr); gap: 12px; margin: 16px 0; }
.stat { background: #1e1e2e; border: 1px solid #313147; border-radius: 10px; padding: 16px; text-align: center; }
.stat-val { font-size: 1.6rem; font-weight: 700; color: #a8edea; }
.stat-lbl { font-size: 0.75rem; color: #888; }
table { width: 100%; border-collapse: collapse; margin: 12px 0; }
th, td { border: 1px solid #313147; padding: 8px 12px; text-align: left; }
th { background: #1e1e2e; color: #a8edea; }
hr { border: none; border-top: 1px solid #313147; margin: 24px 0; }
</style></head><body>
"""
    html += "<h1>🔬 AI Powered Document Q&A System Report</h1>"
    html += f"<p style='text-align:center;color:#888;'>Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p><hr>"

    if stats:
        html += "<h2>📊 Session Analytics</h2><div class='stat-grid'>"
        for key, label in [
            ("questions_asked", "Questions"),
            ("topics_accessed", "Topics"),
            ("sources_used", "Sources"),
            ("total_topics", "Total Topics"),
        ]:
            html += f"<div class='stat'><div class='stat-val'>{stats.get(key, 0)}</div><div class='stat-lbl'>{label}</div></div>"
        html += "</div>"

    if overview:
        html += (
            f"<div class='msg'>{overview}</div>"
        )

    html += "<h2>💬 Research Session</h2>"
    for msg in history:
        role_class = "user" if msg["role"] == "user" else "assistant"
        role_label = (
            "👤 Researcher" if msg["role"] == "user" else "🧠 AI Analyst"
        )
        html += f"<div class='msg {role_class}'><div class='role'>{role_label}</div>"
        if msg.get("intent"):
            html += f"<span class='intent-tag'>{msg['intent'].get('emoji', '')} {msg['intent'].get('label', '')}</span>"
        content = msg["content"].replace("\n", "<br>")
        html += f"<div>{content}</div></div>"

    html += "</body></html>"
    return html


def generate_docx_export(history, overview="", suggestions=None, stats=None):
    doc = DocxDocument()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    title = doc.add_heading(
        "AI Powered Document Q&A System Report",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    if stats:
        doc.add_heading("Session Analytics", level=1)
        table = doc.add_table(rows=1, cols=4, style="Light Grid Accent 1")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        hdr[0].text = "Questions Asked"
        hdr[1].text = "Topics Explored"
        hdr[2].text = "Sources Used"
        hdr[3].text = "Total Topics"
        row = table.add_row().cells
        row[0].text = str(stats.get("questions_asked", 0))
        row[1].text = str(stats.get("topics_accessed", 0))
        row[2].text = str(stats.get("sources_used", 0))
        row[3].text = str(stats.get("total_topics", 0))
        for cell in row:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(14)
        doc.add_paragraph()

    if overview:
        for line in overview.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("## "):
                doc.add_heading(
                    line.replace("## ", "").replace("#", ""), level=2
                )
            elif line.startswith("### "):
                doc.add_heading(
                    line.replace("### ", "").replace("#", ""), level=3
                )
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            else:
                clean = (
                    line.replace("**", "").replace("*", "").replace("`", "")
                )
                doc.add_paragraph(clean)

    if suggestions:
        doc.add_heading("AI-Generated Suggestions", level=1)
        for i, s in enumerate(suggestions, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {s.get('title', f'Suggestion {i}')}")
            run.font.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
            doc.add_paragraph(s.get("description", ""))
            cat = doc.add_paragraph()
            cat_run = cat.add_run(
                f"Category: {s.get('category', 'general').upper()}"
            )
            cat_run.font.italic = True
            cat_run.font.size = Pt(9)
            cat_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_heading("Research Q&A Session", level=1)
    for msg in history:
        role = "Researcher (User)" if msg["role"] == "user" else "AI Analyst"
        p = doc.add_paragraph()
        role_run = p.add_run(f"{role}")
        role_run.font.bold = True
        role_run.font.size = Pt(11)
        if msg["role"] == "user":
            role_run.font.color.rgb = RGBColor(0x1A, 0x5A, 0x8C)
        else:
            role_run.font.color.rgb = RGBColor(0x8C, 0x1A, 0x5A)

        if msg.get("intent") and msg["role"] == "assistant":
            intent = msg["intent"]
            intent_p = doc.add_paragraph()
            intent_run = intent_p.add_run(
                f"[{intent.get('emoji', '')} {intent.get('label', '')}]"
            )
            intent_run.font.italic = True
            intent_run.font.size = Pt(9)
            intent_run.font.color.rgb = RGBColor(0x66, 0x66, 0x99)

        content = msg["content"]
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("## "):
                doc.add_heading(
                    line.replace("## ", "").replace("#", ""), level=3
                )
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            else:
                clean = (
                    line.replace("**", "").replace("*", "").replace("`", "")
                )
                doc.add_paragraph(clean)

        if msg.get("sources") and msg["role"] == "assistant":
            src_p = doc.add_paragraph()
            src_run = src_p.add_run("Sources: ")
            src_run.font.bold = True
            src_run.font.size = Pt(9)
            for src in msg["sources"]:
                src_run = src_p.add_run(f"[{src['file']} | {src['topic']}] ")
                src_run.font.size = Pt(9)
                src_run.font.color.rgb = RGBColor(0x44, 0x44, 0x88)

        doc.add_paragraph("─" * 60).runs[0].font.color.rgb = RGBColor(
            0xCC, 0xCC, 0xCC
        )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_ieee_docx(content, metadata):
    doc = DocxDocument()

    # IEEE papers usually have a specific style.
    # We'll simulate a professional academic layout.

    # Title
    title_text = metadata.get("title") or "Research Paper Title"
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(title_text.upper())
    run_title.font.bold = True
    run_title.font.size = Pt(24)
    run_title.font.name = "Times New Roman"

    # Authors
    p_authors = doc.add_paragraph()
    p_authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_authors = p_authors.add_run(
        metadata.get("authors") or "Author Names Not Provided"
    )
    run_authors.font.size = Pt(11)
    run_authors.font.name = "Times New Roman"

    # Affiliations & Emails
    p_affil = doc.add_paragraph()
    p_affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affil_text = (
        f"{metadata.get('colleges', '')}\n{metadata.get('emails', '')}"
    )
    run_affil = p_affil.add_run(affil_text)
    run_affil.font.italic = True
    run_affil.font.size = Pt(10)
    run_affil.font.name = "Times New Roman"

    doc.add_paragraph()  # Spacer

    # Split content into sections
    sections = content.split("###")
    for section in sections:
        section = section.strip()
        if not section:
            continue

        lines = section.split("\n")
        header = lines[0].strip()
        body = "\n".join(lines[1:]).strip()

        # Section Heading
        p_head = doc.add_paragraph()
        run_head = p_head.add_run(header.upper())
        run_head.font.bold = True
        run_head.font.size = Pt(12)
        run_head.font.name = "Times New Roman"

        # Section Body
        for b_line in body.split("\n"):
            b_line = b_line.strip()
            if not b_line:
                continue

            p_body = doc.add_paragraph()
            p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Remove markdown bold/italics for the DOCX
            clean_line = (
                b_line.replace("**", "").replace("*", "").replace("`", "")
            )

            # Simple bullet point detection
            if b_line.startswith("- ") or b_line.startswith("* "):
                p_body.style = "List Bullet"
                run_body = p_body.add_run(clean_line[2:])
            else:
                run_body = p_body.add_run(clean_line)

            run_body.font.size = Pt(10)
            run_body.font.name = "Times New Roman"

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ─── SIDEBAR ───

with st.sidebar:
    render_ollama_status_sidebar()
    st.markdown("### 📁 Document Upload")
    uploaded_files = st.file_uploader(
        "Select files to process",
        accept_multiple_files=True,
        type=["pdf", "docx", "txt", "csv", "xlsx", "png", "jpg", "jpeg"],
    )

    if st.button("🚀 Process Documents", use_container_width=True):
        if not check_rate_limit("upload", max_requests=5, window_seconds=60):
            st.error("⏳ Rate limit exceeded. Please wait before uploading more documents.")
        elif uploaded_files:
            # Check for concurrent processing lock
            if st.session_state.processing_lock:
                st.warning(
                    "⚠️ Processing in progress. Please wait for the current operation to complete."
                )
            else:
                # Acquire lock
                st.session_state.processing_lock = True
                try:
                    # Sanitize filenames to prevent path traversal attacks
                    data = {
                        sanitize_filename(f.name): process_file(f)[1]
                        for f in uploaded_files
                    }
                    st.session_state.topics_found = {}
                    st.session_state.chat_history = []
                    st.session_state.auto_suggestions = []
                    st.session_state.doc_overview = ""

                    # Reset IEEE metadata conditionally or keep it?
                    # Usually keep it as users might upload new files for the same paper

                    status_text = st.empty()

                    def update_ui(fn, cur, total, topic):
                        if fn not in st.session_state.topics_found:
                            st.session_state.topics_found[fn] = []
                        st.session_state.topics_found[fn].append(topic)
                        status_text.text(
                            f"Processing: {fn}\nIdentified: {topic}"
                        )

                    with st.spinner(
                        "Extracting text and running semantic segmentation..."
                    ):
                        try:
                            st.session_state.ai_engine.ingest_and_segment(
                                data, progress_callback=update_ui
                            )
                        except Exception as e:
                            st.error(f"Error processing documents: {str(e)}")
                            st.session_state.processing_lock = False
                            st.rerun()

                    with st.spinner("Generating document insights..."):
                        try:
                            chunks, metas = (
                                st.session_state.ai_engine.get_all_chunks()
                            )
                            if chunks:
                                st.session_state.doc_overview = st.session_state.ai_engine.research_engine.generate_document_overview(
                                    chunks, metas
                                )
                                st.session_state.auto_suggestions = st.session_state.ai_engine.research_engine.generate_auto_suggestions(
                                    chunks, metas
                                )
                        except Exception as e:
                            st.warning(
                                f"Could not generate insights: {str(e)}"
                            )

                    st.session_state.system_ready = True
                    st.success("✅ Processing complete!")
                except Exception as e:
                    st.error(f"Error processing documents: {str(e)}")
                finally:
                    # Release lock
                    st.session_state.processing_lock = False
        else:
            st.warning("Please upload at least one file.")

    st.divider()

    if st.session_state.chat_history:
        st.divider()
        st.markdown("### 📥 Export Session")

        stats = st.session_state.ai_engine.get_session_stats()
        col1, col2, col3 = st.columns(3)

        with col1:
            st.download_button(
                label="📝 Markdown",
                data=generate_markdown_export(
                    st.session_state.chat_history,
                    st.session_state.doc_overview,
                    st.session_state.auto_suggestions,
                    stats,
                ),
                file_name=f"research_report_{datetime.now().strftime('%Y%m%d_%H%M')}.md",
                mime="text/markdown",
                use_container_width=True,
            )

        with col2:
            st.download_button(
                label="🌐 HTML",
                data=generate_html_export(
                    st.session_state.chat_history,
                    st.session_state.doc_overview,
                    st.session_state.auto_suggestions,
                    stats,
                ),
                file_name=f"research_report_{datetime.now().strftime('%Y%m%d_%H%M')}.html",
                mime="text/html",
                use_container_width=True,
            )

        with col3:
            st.download_button(
                label="📄 Word",
                data=generate_docx_export(
                    st.session_state.chat_history,
                    st.session_state.doc_overview,
                    st.session_state.auto_suggestions,
                    stats,
                ),
                file_name=f"research_report_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
    
    if not st.session_state.get("system_ready", False):
        st.markdown(
            """
        <div style="text-align: center; padding: 30px 15px; border: 1px dashed rgba(99,102,241,0.35); border-radius: 12px; background: rgba(255,255,255,0.85); margin-top: 15px;">
            <h3 style="color: #3730a3 !important; margin: 0; font-size: 1.05rem; font-weight: 700; font-family: 'Inter', sans-serif;">AI Document Q&A</h3>
            <p style="color: #475569 !important; font-size: 0.78rem; margin: 4px 0 0 0; font-weight: 500; font-family: 'Inter', sans-serif;">Document Research & Voice Lab</p>
        </div>
        """,
            unsafe_allow_html=True,
        )


    if (
        st.session_state.system_ready
    ):  # Changed from uploaded_file to system_ready to match existing logic
        st.divider()
        if st.button("🗑️ Clear Session", use_container_width=True):
            st.session_state.chat_history = []
            st.session_state.topics_found = {}
            st.session_state.system_ready = False
            st.session_state.auto_suggestions = []
            st.session_state.doc_overview = ""
            st.rerun()


# ─── MAIN CONTENT ───


if not st.session_state.system_ready:
    st.markdown(
        """
    <div style="text-align: center; padding: 60px 20px;">
        <h2 style="color: #1e1b4b; font-weight: 800; font-size: 2.2rem; background: linear-gradient(90deg, #4f46e5, #7c3aed); -webkit-background-clip: text; -webkit-text-fill-color: transparent;">Welcome to the Research Lab 🧪</h2>
        <p style="color: #475569; font-size: 1.1rem; max-width: 600px; margin: 10px auto 30px auto; line-height: 1.6; font-weight: 500;">
            Upload your documents in the sidebar to begin. The system will automatically
            segment topics, generate insights, and prepare for your research questions.
        </p>
        <div style="margin-top: 30px; display: grid; grid-template-columns: repeat(3, 1fr); gap: 20px; max-width: 800px; margin-left: auto; margin-right: auto;">
            <div style="background: #ffffff; border: 1px solid #e2e8f0; border-radius: 16px; padding: 24px; text-align: center; box-shadow: 0 4px 15px rgba(0, 0, 0, 0.02); transition: transform 0.2s ease;">
                <div style="font-size: 2.2rem; margin-bottom: 10px;">📄</div>
                <div style="color: #0369a1; font-size: 0.95rem; font-weight: 700; margin-top: 8px; text-transform: uppercase; letter-spacing: 0.05em;">Document Q&A</div>
                <div style="color: #64748b; font-size: 0.8rem; margin-top: 6px; font-weight: 500;">Ask factual queries & explore contents</div>
            </div>
            <div style="background: #ffffff; border: 1px solid #e2e8f0; border-radius: 16px; padding: 24px; text-align: center; box-shadow: 0 4px 15px rgba(0, 0, 0, 0.02); transition: transform 0.2s ease;">
                <div style="font-size: 2.2rem; margin-bottom: 10px;">🔬</div>
                <div style="color: #15803d; font-size: 0.95rem; font-weight: 700; margin-top: 8px; text-transform: uppercase; letter-spacing: 0.05em;">Research</div>
                <div style="color: #64748b; font-size: 0.8rem; margin-top: 6px; font-weight: 500;">Propose technical additions & extensions</div>
            </div>
            <div style="background: #ffffff; border: 1px solid #e2e8f0; border-radius: 16px; padding: 24px; text-align: center; box-shadow: 0 4px 15px rgba(0, 0, 0, 0.02); transition: transform 0.2s ease;">
                <div style="font-size: 2.2rem; margin-bottom: 10px;">📥</div>
                <div style="color: #b91c1c; font-size: 0.95rem; font-weight: 700; margin-top: 8px; text-transform: uppercase; letter-spacing: 0.05em;">Export</div>
                <div style="color: #64748b; font-size: 0.8rem; margin-top: 6px; font-weight: 500;">Download complete session report files</div>
            </div>
        </div>
    </div>
    """,
        unsafe_allow_html=True,
    )
else:
    # ─── Chat Tab & Document Overview Tab ───
    # --- Main Navigation ---
    tab_labels = ["💬 Chat", "📄 Overview"]
    active_tab = st.radio("Navigation", tab_labels, horizontal=True, label_visibility="collapsed")
    st.divider()

    if active_tab == "📄 Overview":
        if st.session_state.doc_overview:
            # Render with justify alignment using the new class
            overview_display = strip_image_prompts(st.session_state.doc_overview)
            st.markdown(
                f"""
                <div class="justified-text">
                {overview_display}
                </div>
                """,
                unsafe_allow_html=True
            )
            speak_text(st.session_state.doc_overview, "overview")
        else:
            st.info("No overview generated yet.")
        user_q = None

    elif active_tab == "💬 Chat":
        # Render existing history
        for msg in st.session_state.chat_history:
            with st.chat_message(msg["role"]):
                if msg.get("intent") and msg["role"] == "assistant":
                    intent = msg["intent"]
                    intent_class = f"intent-{intent.get('intent', 'document_qa')}"
                    st.markdown(f'<div class="intent-badge {intent_class}">{intent.get("emoji", "📄")} {intent.get("label", "Document Q&A")}</div>', unsafe_allow_html=True)
                
                if msg.get("reasoning_details") and msg["role"] == "assistant":
                    with st.expander("🧠 View Logic Trace"):
                        st.text(msg["reasoning_details"])

                display_content = strip_image_prompts(msg["content"])
                msg_intent = msg.get("intent", {}).get("intent") if isinstance(msg.get("intent"), dict) else msg.get("intent")
                if msg_intent != "off_topic":
                    render_content_with_mermaid(display_content)
                else:
                    st.markdown(display_content)

                if msg["role"] == "assistant":
                    speak_text(msg["content"], f"hist_{id(msg)}")

                if msg.get("confidence") is not None and msg["role"] == "assistant":
                    render_confidence_badge(msg["confidence"])

                if msg.get("sources") and msg["role"] == "assistant":
                    with st.expander(f"📚 Sources Referenced ({len(msg['sources'])})"):
                        for src in msg["sources"]:
                            st.markdown(f"""
                            <div class="source-card">
                                <div class="source-header">
                                    <span class="source-file">{sanitize_for_markdown(src.get('file', ''))}</span>
                                    <span class="source-score">{src['score']}</span>
                                </div>
                                <div class="source-topic">{sanitize_for_markdown(src.get('topic', ''))}</div>
                                <div class="source-preview">{sanitize_for_markdown(src.get('preview', ''))}</div>
                            </div>
                            """, unsafe_allow_html=True)

                if msg["role"] == "assistant" and msg.get("intent", {}).get("intent") == "ieee_paper_gen":
                    ieee_docx = generate_ieee_docx(msg["content"], st.session_state.ieee_metadata)
                    st.download_button(label="📄 Download IEEE Official Paper (.docx)", data=ieee_docx, file_name=f"IEEE_Paper.docx", key=f"dl_ieee_{id(msg)}", use_container_width=True)

        # --- Question bar below the answers with mic inside the bar ---
        mic_text = st.session_state.pop("mic_text", None)
        
        # Chat input — Streamlit pins this to bottom automatically
        user_q = st.chat_input(
            "Ask about your documents, request suggestions, or propose research add-ons...",
            max_chars=MAX_CHAT_INPUT_LENGTH,
            key="chat_input_main"
        )
        # Mic button overlaid on the chat input bar via fixed CSS positioning
        st.markdown('<div class="mic-overlay-btn">', unsafe_allow_html=True)
        render_mic_input()
        st.markdown('</div>', unsafe_allow_html=True)
            
        if mic_text and not user_q:
            user_q = mic_text
    else:
        user_q = None

    if user_q:
        if not check_rate_limit("chat", max_requests=10, window_seconds=60):
            st.error("⏳ Rate limit exceeded. Please wait before sending more messages.")
            st.stop()
        if len(user_q) > MAX_CHAT_INPUT_LENGTH:
            st.error(
                f"Input too long. Maximum {MAX_CHAT_INPUT_LENGTH} characters allowed."
            )
            st.rerun()
        if not user_q.strip():
            st.warning("Please enter a valid question.")
            st.rerun()

        st.session_state.chat_history.append(
            {"role": "user", "content": user_q}
        )

        with st.chat_message("user"):
            st.markdown(user_q)

        with st.chat_message("assistant"):
            stream_meta = {}
            final_data = {}
            confidence_result = None

            response_container = st.empty()
            streamed_text = ""

            for event in st.session_state.ai_engine.get_answer_stream(
                user_q,
                st.session_state.chat_history[:-1],
                metadata=st.session_state.ieee_metadata,
            ):
                if event["type"] == "meta":
                    stream_meta = event
                    intent = event.get("intent", {})
                    intent_class = (
                        f"intent-{intent.get('intent', 'document_qa')}"
                    )
                    st.markdown(
                        f'<div class="intent-badge {intent_class}">{intent.get("emoji", "📄")} {intent.get("label", "Document Q&A")}</div>',
                        unsafe_allow_html=True,
                    )
                    confidence_result = event.get("confidence")
                elif event["type"] == "token":
                    streamed_text += event["token"]
                    response_container.markdown(streamed_text + "▌")
                elif event["type"] == "confidence_update":
                    confidence_result = event.get("confidence", confidence_result)
                elif event["type"] == "done":
                    final_data = event

            final_content = final_data.get("content", streamed_text)
            response_container.empty()

            # Strip image prompts from displayed content
            display_final = strip_image_prompts(final_content)
            if stream_meta.get("intent", {}).get("intent") != "off_topic":
                render_content_with_mermaid(display_final)
            else:
                st.markdown(display_final)

            speak_text(final_content, "stream_latest")

            reasoning = final_data.get("reasoning")
            if reasoning:
                with st.expander("🧠 View Logic Trace"):
                    st.text(reasoning)

            # Confidence badge — rendered AFTER answer content
            if confidence_result is not None:
                render_confidence_badge(confidence_result)

            sources = stream_meta.get("sources", [])
            if sources:
                with st.expander(
                    f"📚 Sources Referenced ({len(sources)})"
                ):
                    for src in sources:
                        safe_file = sanitize_for_markdown(
                            src.get("file", "")
                        )
                        safe_topic = sanitize_for_markdown(
                            src.get("topic", "")
                        )
                        safe_preview = sanitize_for_markdown(
                            src.get("preview", "")
                        )
                        st.markdown(
                            f"""
                        <div class="source-card">
                            <div class="source-header">
                                <span class="source-file">{safe_file}</span>
                                <span class="source-score">{src['score']}</span>
                            </div>
                            <div class="source-topic">{safe_topic}</div>
                            <div class="source-preview">{safe_preview}</div>
                        </div>
                        """,
                            unsafe_allow_html=True,
                        )

            if (
                stream_meta.get("intent", {}).get("intent")
                == "ieee_paper_gen"
            ):
                ieee_docx = generate_ieee_docx(
                    final_content, st.session_state.ieee_metadata
                )
                st.download_button(
                    label="📄 Download IEEE Official Paper (.docx)",
                    data=ieee_docx,
                    file_name=f"IEEE_Paper_{st.session_state.ieee_metadata.get('title', 'Generated').replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_ieee_latest",
                    use_container_width=True,
                )

        st.session_state.chat_history.append(
            {
                "role": "assistant",
                "content": final_content,
                "reasoning_details": final_data.get("reasoning"),
                "intent": stream_meta.get("intent"),
                "sources": stream_meta.get("sources", []),
                "confidence": confidence_result,
            }
        )
